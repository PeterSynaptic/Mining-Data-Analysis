import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import os
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Mining Data Analysis", layout="wide")

# Title and description
st.title("Mining Data Analysis Dashboard")
st.markdown("Upload your mining dataset to generate analysis and visualizations.")

# File uploader
uploaded_file = st.file_uploader("Choose an Excel file", type=['xlsx'])

def generate_report_doc():
    doc = Document()
    doc.add_heading('Mining Data Analysis Report', 0)
    
    # Add summary statistics
    doc.add_heading('Summary Statistics', level=1)
    for key, value in summary.items():
        doc.add_paragraph(f"{key}: {value}")
    
    # Add visualization descriptions
    doc.add_heading('Visualization Analysis', level=1)
    
    doc.add_heading('1. Grade Distributions', level=2)
    doc.add_paragraph('The plots show the distribution of copper and molybdenum grades in the dataset, '
                     'helping identify the typical concentration ranges and any potential anomalies.')
    
    doc.add_heading('2. Prediction Accuracy', level=2)
    doc.add_paragraph('These plots compare sensor predictions with actual blasthole assays, '
                     'demonstrating the accuracy of the sensor measurements.')
    
    doc.add_heading('3. Distance vs Prediction Error', level=2)
    doc.add_paragraph('This analysis shows how the distance to the nearest blasthole affects '
                     'the accuracy of grade predictions.')
    
    doc.add_heading('4. Shift Analysis', level=2)
    doc.add_paragraph('The plot compares average grades across different shifts, '
                     'helping identify any systematic variations in measurements.')
    
    return doc

if uploaded_file is not None:
    # Read the data
    df = pd.read_excel(uploaded_file)
    
    # Create plots directory if it doesn't exist
    if not os.path.exists('plots'):
        os.makedirs('plots')
    
    # Display basic dataset information
    st.subheader("Dataset Overview")
    st.write(f"Number of records: {len(df)}")
    st.write(f"Number of columns: {len(df.columns)}")
    
    # Display sample of the data
    st.subheader("Sample Data")
    st.dataframe(df.head())
    
    # Create visualizations
    st.subheader("Data Visualizations")
    
    # 1. Grade Distribution Analysis
    st.write("### Grade Distributions")
    fig1, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    sns.histplot(data=df, x='cugrade', bins=30, ax=ax1)
    ax1.set_title('Distribution of Copper Grade Predictions')
    ax1.set_xlabel('Copper Grade (%)')
    
    sns.histplot(data=df, x='mograde', bins=30, ax=ax2)
    ax2.set_title('Distribution of Molybdenum Grade Predictions')
    ax2.set_xlabel('Molybdenum Grade (%)')
    plt.tight_layout()
    st.pyplot(fig1)
    
    # 2. Prediction vs Actual Analysis
    st.write("### Prediction Accuracy")
    fig2, (ax3, ax4) = plt.subplots(1, 2, figsize=(12, 5))
    ax3.scatter(df['cugrade'], df['avg_bh_grade_cu'], alpha=0.5)
    ax3.plot([df['cugrade'].min(), df['cugrade'].max()], 
             [df['cugrade'].min(), df['cugrade'].max()], 'r--')
    ax3.set_title('Sensor Cu Prediction vs Blasthole Cu Grade')
    ax3.set_xlabel('Sensor Predicted Cu Grade (%)')
    ax3.set_ylabel('Blasthole Cu Grade (%)')
    
    ax4.scatter(df['mograde'], df['avg_bh_grade_mo'], alpha=0.5)
    ax4.plot([df['mograde'].min(), df['mograde'].max()], 
             [df['mograde'].min(), df['mograde'].max()], 'r--')
    ax4.set_title('Sensor Mo Prediction vs Blasthole Mo Grade')
    ax4.set_xlabel('Sensor Predicted Mo Grade (%)')
    ax4.set_ylabel('Blasthole Mo Grade (%)')
    plt.tight_layout()
    st.pyplot(fig2)
    
    # 3. Distance vs Prediction Error Analysis
    st.write("### Distance vs Prediction Error")
    df['cu_prediction_error'] = abs(df['cugrade'] - df['avg_bh_grade_cu'])
    fig3, ax5 = plt.subplots(figsize=(10, 6))
    ax5.scatter(df['Dist_to_NN_bh'], df['cu_prediction_error'], alpha=0.5)
    ax5.set_title('Distance to Nearest Blasthole vs Cu Prediction Error')
    ax5.set_xlabel('Distance to Nearest Blasthole (m)')
    ax5.set_ylabel('Absolute Cu Prediction Error (%)')
    plt.tight_layout()
    st.pyplot(fig3)
    
    # 4. Shift Analysis
    st.write("### Shift Analysis")
    fig4, ax6 = plt.subplots(figsize=(14, 8))
    
    # Group by unique shift_id and take every 5th shift
    shift_stats = df.groupby('shift_id')[['cugrade', 'mograde']].mean().sort_index()
    shift_stats = shift_stats.iloc[::5]  # Take every 5th shift
    
    # Shorten shift IDs for display (keep only the date part)
    short_labels = [sid.split('D')[0] for sid in shift_stats.index]
    shift_stats.index = short_labels
    
    # Use distinct colors for better visibility
    bars = shift_stats.plot(kind='bar', ax=ax6, width=0.8, color=['#2ecc71', '#e74c3c'])
    ax6.set_title('Average Grades by Shift Date', fontsize=12, pad=20)
    ax6.set_xlabel('Shift Date', fontsize=10)
    ax6.set_ylabel('Grade (%)', fontsize=10)
    ax6.legend(['Copper', 'Molybdenum'], fontsize=10, loc='upper right')
    ax6.grid(True, axis='y', linestyle='--', alpha=0.7)
    
    # Improve x-axis labels
    plt.xticks(rotation=30, ha='right')
    
    # Add value labels on top of bars with better contrast
    for container in bars.containers:
        ax6.bar_label(container, fmt='%.2f', padding=3, fontsize=9)
    
    # Adjust layout
    plt.subplots_adjust(bottom=0.2)
    plt.tight_layout()
    st.pyplot(fig4)
    
    # Generate summary statistics
    summary = {
        'Average Copper Grade': f"{df['cugrade'].mean():.3f}%",
        'Average Molybdenum Grade': f"{df['mograde'].mean():.3f}%",
        'Cu Grade Range': f"{df['cugrade'].min():.3f}% - {df['cugrade'].max():.3f}%",
        'Mo Grade Range': f"{df['mograde'].min():.3f}% - {df['mograde'].max():.3f}%",
        'Average Distance to Nearest Blasthole': f"{df['Dist_to_NN_bh'].mean():.2f}m",
        'Total Number of Measurements': len(df),
        'Date Range': f"{df['run_date_time'].min()} to {df['run_date_time'].max()}"
    }
    
    # Display summary statistics
    st.subheader("Summary Statistics")
    for key, value in summary.items():
        st.write(f"**{key}:** {value}")
    
    # Generate and download report
    st.subheader("Download Report")
    if st.button("Generate Report"):
        doc = generate_report_doc()
        # Save the report to BytesIO object
        bio = BytesIO()
        doc.save(bio)
        
        # Download button for the report
        st.download_button(
            label="Download Report as Word Document",
            data=bio.getvalue(),
            file_name="mining_analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Please upload an Excel file to begin the analysis.")